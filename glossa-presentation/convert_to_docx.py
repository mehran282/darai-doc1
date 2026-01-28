#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Convert Markdown to DOCX using python-docx
"""

import os
import sys
import re

def install_package(package):
    """Try to install package"""
    try:
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", package, "--user"])
        return True
    except:
        return False

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("Installing python-docx...")
    if not install_package("python-docx"):
        print("Could not install python-docx. Please install manually: pip install python-docx")
        sys.exit(1)
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

def convert_to_persian_numbers(text):
    """Convert English digits to Persian digits"""
    persian_digits = '۰۱۲۳۴۵۶۷۸۹'
    english_digits = '0123456789'
    translation_table = str.maketrans(english_digits, persian_digits)
    return text.translate(translation_table)

def is_table_separator(line):
    """Check if line is a markdown table separator"""
    stripped = line.strip()
    if not '|' in stripped:
        return False
    # Check if it's a separator like |---|---| or |:---|:---:|---:|
    if re.match(r'^\|[\s\-:]+\|', stripped):
        return True
    # Check if it contains mostly dashes and pipes
    if stripped.replace('-', '').replace('|', '').replace(':', '').replace(' ', '') == '':
        return True
    return False

def parse_markdown_table(table_text):
    """Parse markdown table"""
    lines = table_text.strip().split('\n')
    rows = []
    for line in lines:
        # Skip separator lines
        if is_table_separator(line):
            continue
        if '|' in line:
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            rows.append(cells)
    return rows

def add_table_to_doc(doc, table_text):
    """Add markdown table to Word document"""
    rows = parse_markdown_table(table_text)
    if not rows:
        return
    
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(rows):
        for j, cell_data in enumerate(row_data):
            # Remove markdown formatting
            cell_data = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_data)
            cell_data = re.sub(r'`(.*?)`', r'\1', cell_data)
            cell_data = re.sub(r'&nbsp;', ' ', cell_data)
            cell = table.rows[i].cells[j]
            # Convert numbers to Persian
            cell_data = convert_to_persian_numbers(cell_data)
            # Fix parentheses for RTL text
            cell_data = cell_data.replace('(', '\u200E(').replace(')', ')\u200E')
            cell.text = cell_data
            # Set RTL for Persian text
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # Set RTL direction for the paragraph
                paragraph_format = paragraph.paragraph_format
                paragraph_format.right_to_left = True

def convert_md_to_docx(md_file, docx_file):
    """Convert Markdown file to Word"""
    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        doc = Document()
        
        # Set RTL direction and page size
        sections = doc.sections
        for section in sections:
            section.page_height = Inches(11.69)
            section.page_width = Inches(8.27)
            # Set RTL for the section
            try:
                section._sectPr.xpath('./w:bidi')[0].set(qn('w:val'), '1')
            except:
                pass
        
        lines = content.split('\n')
        i = 0
        table_lines = []
        in_table = False
        in_code_block = False
        
        while i < len(lines):
            line = lines[i]
            
            # Code block
            if line.strip().startswith('```'):
                in_code_block = not in_code_block
                i += 1
                continue
            
            if in_code_block:
                i += 1
                continue
            
            # Skip table separator lines completely
            if is_table_separator(line):
                # If we're in a table, separator is just formatting - skip it
                if in_table:
                    i += 1
                    continue
                # If not in table, separator doesn't make sense - skip it
                i += 1
                continue
            
            # Check if it's a table row
            if '|' in line:
                if not in_table:
                    in_table = True
                    table_lines = []
                table_lines.append(line)
            elif in_table:
                # End of table - process it
                if table_lines:
                    add_table_to_doc(doc, '\n'.join(table_lines))
                    table_lines = []
                in_table = False
            
            if not in_table:
                # Headers
                if line.startswith('# '):
                    text = line[2:].strip()
                    text = convert_to_persian_numbers(text)
                    text = text.replace('(', '\u200E(').replace(')', ')\u200E')
                    p = doc.add_heading(text, level=1)
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p.paragraph_format.right_to_left = True
                elif line.startswith('## '):
                    text = line[3:].strip()
                    text = convert_to_persian_numbers(text)
                    text = text.replace('(', '\u200E(').replace(')', ')\u200E')
                    p = doc.add_heading(text, level=2)
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p.paragraph_format.right_to_left = True
                elif line.startswith('### '):
                    text = line[4:].strip()
                    text = convert_to_persian_numbers(text)
                    text = text.replace('(', '\u200E(').replace(')', ')\u200E')
                    p = doc.add_heading(text, level=3)
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p.paragraph_format.right_to_left = True
                elif line.startswith('#### '):
                    text = line[5:].strip()
                    text = convert_to_persian_numbers(text)
                    text = text.replace('(', '\u200E(').replace(')', ')\u200E')
                    p = doc.add_heading(text, level=4)
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p.paragraph_format.right_to_left = True
                # Horizontal rule
                elif line.strip() == '---':
                    p = doc.add_paragraph('─' * 50)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # List items
                elif line.strip().startswith('- ') or line.strip().startswith('* '):
                    text = line.strip()[2:]
                    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                    text = convert_to_persian_numbers(text)
                    text = text.replace('(', '\u200E(').replace(')', ')\u200E')
                    p = doc.add_paragraph(text, style='List Bullet')
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p.paragraph_format.right_to_left = True
                # Regular paragraph
                elif line.strip() and not line.strip().startswith('**تهیه شده'):
                    # Remove markdown formatting
                    text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                    text = re.sub(r'`(.*?)`', r'\1', text)
                    # Convert numbers to Persian
                    text = convert_to_persian_numbers(text)
                    # Fix parentheses for RTL text - add LTR mark before ( and after )
                    text = text.replace('(', '\u200E(').replace(')', ')\u200E')
                    p = doc.add_paragraph(text.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    # Set RTL direction
                    p.paragraph_format.right_to_left = True
            
            i += 1
        
        # Final table
        if table_lines:
            add_table_to_doc(doc, '\n'.join(table_lines))
        
        # Set font and RTL for all paragraphs
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.right_to_left = True
            for run in paragraph.runs:
                run.font.name = 'Vazirmatn'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Vazirmatn')
        
        # Set RTL for all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.right_to_left = True
        
        doc.save(docx_file)
        print("DOCX file created successfully!")
        return True
    except Exception as e:
        print("Error occurred during conversion")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    md_file = "برآورد-صرفه-جویی-جایگزینی-نیروی-انسانی.md"
    docx_file = "برآورد-صرفه-جویی-جایگزینی-نیروی-انسانی-اصلاح-شده.docx"
    
    if not os.path.exists(md_file):
        print("Markdown file not found!")
        sys.exit(1)
    
    print("Converting markdown to DOCX...")
    convert_md_to_docx(md_file, docx_file)
