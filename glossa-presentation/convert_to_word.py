#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script to convert Markdown to Word
"""

import sys
import os
import re

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import markdown
except ImportError:
    print("Installing required packages...")
    os.system("pip install python-docx markdown")
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import markdown

def parse_markdown_table(table_text):
    """Parse markdown table and return rows"""
    lines = table_text.strip().split('\n')
    rows = []
    for line in lines:
        if '|' in line and not line.strip().startswith('|---'):
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            rows.append(cells)
    return rows

def add_table_to_doc(doc, table_text):
    """Add markdown table to Word document"""
    rows = parse_markdown_table(table_text)
    if not rows:
        return
    
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(rows):
        for j, cell_data in enumerate(row_data):
            # Remove markdown formatting
            cell_data = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_data)
            cell_data = re.sub(r'`(.*?)`', r'\1', cell_data)
            table.rows[i].cells[j].text = cell_data

def convert_md_to_docx(md_file, docx_file):
    """Convert Markdown file to Word"""
    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        doc = Document()
        
        # Set RTL direction (if needed)
        sections = doc.sections
        for section in sections:
            section.page_height = Inches(11.69)
            section.page_width = Inches(8.27)
        
        lines = content.split('\n')
        i = 0
        table_lines = []
        in_table = False
        
        while i < len(lines):
            line = lines[i]
            
            # Check if it's a table
            if '|' in line and not line.strip().startswith('---'):
                if not in_table:
                    in_table = True
                    table_lines = []
                table_lines.append(line)
            elif in_table:
                # End of table
                if table_lines:
                    add_table_to_doc(doc, '\n'.join(table_lines))
                    table_lines = []
                in_table = False
            
            if not in_table:
                # Headers
                if line.startswith('# '):
                    p = doc.add_heading(line[2:].strip(), level=1)
                elif line.startswith('## '):
                    p = doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith('### '):
                    p = doc.add_heading(line[4:].strip(), level=3)
                elif line.startswith('#### '):
                    p = doc.add_heading(line[5:].strip(), level=4)
                # Horizontal rule
                elif line.strip() == '---':
                    doc.add_paragraph('─' * 50)
                # Code block
                elif line.strip().startswith('```'):
                    pass  # Skip code block markers
                # Regular paragraph
                elif line.strip() and not line.strip().startswith('**تهیه شده'):
                    # Remove markdown formatting
                    text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                    text = re.sub(r'`(.*?)`', r'\1', text)
                    p = doc.add_paragraph(text.strip())
            
            i += 1
        
        # Add final table if exists
        if table_lines:
            add_table_to_doc(doc, '\n'.join(table_lines))
        
        doc.save(docx_file)
        print(f"File {docx_file} created successfully!")
        return True
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    md_file = "برآورد-صرفه-جویی-جایگزینی-نیروی-انسانی.md"
    docx_file = "برآورد-صرفه-جویی-جایگزینی-نیروی-انسانی.docx"
    
    if not os.path.exists(md_file):
        print(f"File {md_file} not found!")
        sys.exit(1)
    
    print(f"Converting {md_file} to {docx_file}...")
    convert_md_to_docx(md_file, docx_file)
